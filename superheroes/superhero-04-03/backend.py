import os
import json
import tempfile
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import whisper
from docx import Document
import google.generativeai as genai

import markdown2
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from google.cloud import secretmanager


app = Flask(__name__)
CORS(app)  # Enable CORS for React frontend

# Load Whisper model once
model = whisper.load_model("base")

# File to store meeting history
HISTORY_FILE = "meeting_history.json"

# Load history from the file
def load_history():
    if os.path.exists(HISTORY_FILE):
        with open(HISTORY_FILE, "r") as file:
            return json.load(file)
    return []

# Get the gemini secret from the environment
def get_secret():
    client = secretmanager.SecretManagerServiceClient()
    name = f"projects/150699885662/secrets/superhero-04-03-secret/versions/latest"
    response = client.access_secret_version(name=name)
    return response.payload.data.decode("UTF-8")


# Save history to the file
def save_history(history):
    with open(HISTORY_FILE, "w") as file:
        json.dump(history, file, indent=4)

# Function to transcribe audio using Whisper
def transcribe_audio(mp3_file_path):
    transcription = model.transcribe(mp3_file_path)
    return transcription['text']

# Function to summarize the transcription
def summarize_transcription(transcription):
    # will not work
    google_api_key = get_secret()
    genai.configure(api_key=google_api_key)

    model = genai.GenerativeModel("gemini-1.5-flash")
    response = model.generate_content(
        f"""Hello Gemini. Take the following text as a transcription from a software development meeting between stakeholders. Agile and Scrum techniques will be applied. 
        - Generate epics, themes, and user stories, each with an estimation number.
        - Additionally, divide the generated requirements into two categories: Functional Requirements and Non-Functional Requirements.
        - Provide a clear and structured list for each category If you can store the user stories in two different markdown tables, divided into functional and non fucntional, that would be awesome. 

        Here is the transcription:
        {transcription}
        """
    )
    return response.text




def add_html_to_docx(doc, html):
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html, "html.parser")

    for element in soup:
        if element.name == "p":
            doc.add_paragraph(element.get_text())
        elif element.name == "ul":
            for li in element.find_all("li"):
                doc.add_paragraph(li.get_text(), style='List Bullet')
        elif element.name == "ol":
            for li in element.find_all("li"):
                doc.add_paragraph(li.get_text(), style='List Number')
        elif element.name == "h1":
            doc.add_heading(element.get_text(), level=1)
        elif element.name == "h2":
            doc.add_heading(element.get_text(), level=2)
        elif element.name == "h3":
            doc.add_heading(element.get_text(), level=3)
        elif element.name == "h4":
            doc.add_heading(element.get_text(), level=4)
        elif element.name == "h5":
            doc.add_heading(element.get_text(), level=5)
        elif element.name == "h6":
            doc.add_heading(element.get_text(), level=6)
        elif element.name == "table":
            table = doc.add_table(rows=0, cols=0)
            for row in element.find_all("tr"):
                cells = row.find_all(["td", "th"])
                if len(cells) > 0:
                    tr = table.add_row().cells
                    for i, cell in enumerate(cells):
                        if len(table.columns) < len(cells):
                            table.add_column()
                        tr[i].text = cell.get_text()
                        if cell.name == "th":
                            tr[i].paragraphs[0].runs[0].bold = True
        



# Function to generate IEEE SRS document
def generate_srs(transcription, summary):
    output_file = "SRS_Document.docx"
    doc = Document()
    doc.add_heading('Software Requirements Specification (SRS)', level=1)

    doc.add_heading('1. Introduction', level=2)
    
    # Add bold text for "Purpose:"
    p = doc.add_paragraph()
    p.add_run("Purpose:").bold = True
    p.add_run(" This document defines the requirements extracted from a transcription of stakeholder meetings.")
    
    # Add bold text for "Scope:"
    p = doc.add_paragraph()
    p.add_run("Scope:").bold = True
    p.add_run(" The system is designed to facilitate transcription and requirement extraction for software development projects using Agile methodologies.")
    
    # Add bold text for "References:"
    p = doc.add_paragraph()
    p.add_run("References:").bold = True
    p.add_run(" IEEE STD 830-1998.")

    doc.add_heading('2. Overall Description', level=2)
    
    # Add bold text for "Product Perspective:"
    p = doc.add_paragraph()
    p.add_run("Product Perspective:").bold = True
    p.add_run(" The system integrates Whisper for transcription and Google Gemini for requirement extraction.")
    
    # Add bold text for "Product Functions:"
    p = doc.add_paragraph()
    p.add_run("Product Functions:").bold = True
    p.add_run(" The primary function is to convert audio meeting transcriptions into actionable requirements (epics, themes, and user stories).")
    
    # Add bold text for "Constraints:"
    p = doc.add_paragraph()
    p.add_run("Constraints:").bold = True
    p.add_run(" The system relies on the Whisper and Google Gemini APIs for functionality.")

    doc.add_heading('3. Specific Requirements', level=2)
    doc.add_paragraph("The requirements extracted are as follows:")

    doc.add_heading('Original Transcription', level=3)
    doc.add_paragraph(transcription)


    doc.add_heading('Extracted Requirements', level=3)
    html_summary = markdown2.markdown(summary)
    add_html_to_docx(doc, html_summary)

    doc.save(output_file)
    return output_file

@app.route("/api/transcribe", methods=["POST"])
def transcribe():

    print("Request received:", request)
    # Get the meeting title and uploaded file
    meeting_title = request.form.get("title")
    mp3_file = request.files.get("audio")

    if not meeting_title or not mp3_file:
        return jsonify({"error": "Meeting title and audio file are required."}), 400

    # Save the uploaded MP3 file to a temporary location
    with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as tmp_file:
        tmp_file.write(mp3_file.read())
        tmp_file_path = tmp_file.name

    # Transcribe and summarize
    transcription = transcribe_audio(tmp_file_path)
    summary = summarize_transcription(transcription)

    # Save to history
    history = load_history()
    new_entry = {
        "title": meeting_title,
        "transcription": transcription,
        "summary": summary,
    }
    history.append(new_entry)
    save_history(history)

    # Clean up temp file
    os.remove(tmp_file_path)

    return jsonify({"transcription": transcription, "summary": summary})

@app.route("/api/download_srs", methods=["POST"])
def download_srs():
    data = request.json
    transcription = data.get("transcription")
    summary = data.get("summary")

    if not transcription or not summary:
        return jsonify({"error": "Transcription and summary are required."}), 400

    # Generate SRS document
    srs_path = generate_srs(transcription, summary)
    return send_file(srs_path, as_attachment=True)

@app.route("/api/history", methods=["GET"])
def history():
    return jsonify(load_history())

if __name__ == "__main__":
    print("Starting Flask app on port 5000...")
    app.run(host="0.0.0.0", port=5000)